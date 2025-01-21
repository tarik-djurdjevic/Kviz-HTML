import json
from docx import Document

file_name = "Uredsko poslovanje u organima uprave-pitanja"
# Load the Word document
doc_path = file_name + ".docx"
doc = Document(doc_path)

# Initialize an empty list to store quiz data
quiz_data = []

# Process tables in the document
for table in doc.tables:
    rows = table.rows
    for i in range(len(rows)):
        cells = rows[i].cells
        print(len(cells))
        if len(cells) > 2:
          # Normalize the text for comparison
          cell_text = cells[1].text.strip().lower()
          if "pitanje" in cell_text:  # Detect question row
            print("YES!!!!!")
            print("pitanje" in cell_text, cell_text)
            question = cells[2].text.strip()
            # print(f"Question found: {question}")  # Debug: Print question
            options = []
            correct_index = None

            # Process the next 4 rows for options
            for j in range(1, 5):
                if i + j < len(rows):
                    option_cells = rows[i + j].cells
                    if len(option_cells) > 2:
                        option_text = option_cells[2].text.strip()
                        # print(f"Option found: {option_text}")  # Debug: Print option
                        underlined = any(run.underline for run in option_cells[2].paragraphs[0].runs)
                        options.append(option_text)

                        # If the option is underlined, capture its index
                        if underlined:
                            # print(f"Underlined option: {option_text}")  # Debug: Print underlined option
                            correct_index = len(options) - 1

            # Add question and options to quiz data
            if correct_index is not None:
                quiz_data.append({
                    "question": question,
                    "options": options,
                    "correct": correct_index
                })

# Output the first two questions
output_data = quiz_data[:2]
# print(json.dumps(output_data, ensure_ascii=False, indent=2).encode('utf-8').decode('utf-8'))

# Specify the file path where you want to save the data
file_path = file_name + '.json'

# Open the file and write the JSON data
with open(file_path, 'w', encoding='utf-8') as f:
    json.dump(quiz_data, f, ensure_ascii=False, indent=2)

print(f"Data saved to json file")